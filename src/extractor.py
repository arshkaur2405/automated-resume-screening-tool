"""
Resume Text Extraction Module

This module handles extracting text content from PDF and DOCX resume files.
Supports multiple file formats commonly used for resumes.
"""

import os
from pathlib import Path
from typing import Optional

# PDF extraction
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX extraction  
try:
    import docx
except ImportError:
    docx = None


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as a single string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If PyPDF2 is not installed
    """
    if PyPDF2 is None:
        raise ValueError("PyPDF2 is not installed. Run: pip install PyPDF2")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    text_content = []
    
    try:
        with open(file_path, 'rb') as file:
            # Create PDF reader object
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
                    
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return ""
    
    # Join all pages with newlines
    return '\n'.join(text_content)


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text content from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text as a single string
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If python-docx is not installed
    """
    if docx is None:
        raise ValueError("python-docx is not installed. Run: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    text_content = []
    
    try:
        # Load the document
        document = docx.Document(file_path)
        
        # Extract text from each paragraph
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
                        
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return ""
    
    return '\n'.join(text_content)


def extract_text_from_txt(file_path: str) -> str:
    """
    Read text content from a plain text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        File content as string
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()


def extract_text(file_path: str) -> str:
    """
    Extract text from any supported file format.
    Automatically detects file type based on extension.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text content
        
    Raises:
        ValueError: If file format is not supported
    """
    file_extension = Path(file_path).suffix.lower()
    
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.doc': extract_text_from_docx,  # May not work for old .doc format
        '.txt': extract_text_from_txt,
    }
    
    if file_extension not in extractors:
        raise ValueError(
            f"Unsupported file format: {file_extension}. "
            f"Supported formats: {list(extractors.keys())}"
        )
    
    return extractors[file_extension](file_path)


def get_candidate_name_from_filename(file_path: str) -> str:
    """
    Extract candidate name from resume filename.
    Assumes format like 'resume_john_doe.pdf' or 'John_Doe_Resume.pdf'
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Formatted candidate name
    """
    # Get filename without extension
    filename = Path(file_path).stem
    
    # Remove common prefixes/suffixes
    for remove_word in ['resume', 'cv', 'curriculum_vitae']:
        filename = filename.lower().replace(remove_word, '')
    
    # Replace underscores and hyphens with spaces
    name = filename.replace('_', ' ').replace('-', ' ')
    
    # Clean up and title case
    name = ' '.join(name.split())  # Remove extra spaces
    name = name.title()
    
    return name if name else "Unknown Candidate"


if __name__ == "__main__":
    # Test the extraction functions
    print("Resume Text Extractor Module")
    print("Supported formats: PDF, DOCX, TXT")
